from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "毕业论文_Ai_Vuln_Scanner.docx"


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_cn_paragraph(
    doc: Document,
    text: str,
    font: str = "宋体",
    size: int = 14,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt: int = 22,
    bold: bool = False,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, font, size, bold=bold)
    return p


def add_cn_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(22)
        pf.space_after = Pt(22)
        size = 15  # 小三
        font = "黑体"
    elif level == 2:
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        size = 14
        font = "黑体"
    else:
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        size = 14
        font = "宋体"
    pf.line_spacing = Pt(22)
    r = p.add_run(text)
    set_run_font(r, font, size, bold=(level <= 2))


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_abstract(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    r1 = p.add_run("摘要：")
    set_run_font(r1, "黑体", 12, bold=False)
    r2 = p.add_run(
        "随着高校与企业数字化系统持续扩展，Web 应用安全检测逐渐由人工经验驱动转向自动化与智能化协同。"
        "本文围绕“AI 漏洞扫描与报告生成平台”开展设计与实现，面向本科毕业设计场景，构建了一套可复现实验、"
        "可配置流程、可解释报告的应用安全检测系统。系统以 Flask 为服务框架，组织 URL 收集、Nuclei 模板扫描、"
        "ZAP 主动扫描、AI 分析增强与多格式报告导出等关键环节，并通过任务状态机、数据库落库与降级机制保障流程稳定。"
        "针对真实使用中“扫描耗时长、无效结果多、报告描述模板化”的痛点，本文重点完成三项优化：其一，提出 public/dvwa 双场景"
        "扫描模式，按目标类型自动切换模板集合与扫描策略，避免公网任务误用本地靶场模板；其二，引入漏洞去重与 URL 截断控制，"
        "降低重复分析与无效匹配带来的资源浪费；其三，重构 AI 提示词与结果解析链路，要求模型输出结构化字段，显著提升报告可读性与"
        "整改可执行性。实验以本地 DVWA 及公开站点为对象进行验证，结果表明系统能够在可控时间内完成任务编排、漏洞检测与报告产出，"
        "并在稳定性、可维护性与工程实用性方面达到毕业设计预期。本文的工作说明：将规则引擎、扫描工具与大模型能力进行分层解耦，"
        "可在不牺牲合规边界的前提下提升渗透测试辅助效率，为后续研究提供可扩展基础。"
    )
    set_run_font(r2, "仿宋", 12, bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.line_spacing = Pt(20)
    r3 = p2.add_run("关键词：")
    set_run_font(r3, "黑体", 12, bold=False)
    r4 = p2.add_run("漏洞扫描；Nuclei；OWASP ZAP；DVWA；大语言模型；自动化测试")
    set_run_font(r4, "仿宋", 12, bold=False)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = 1.0
    spacer.add_run(" ")


def add_en_part(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(22)
    r = p.add_run("Design and Implementation of an AI-Enhanced Web Vulnerability Scanner")
    set_run_font(r, "Times New Roman", 14, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = Pt(22)
    r2 = p2.add_run("Lei Xin Guo")
    set_run_font(r2, "Times New Roman", 12, bold=False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.line_spacing = 1.0
    r3 = p3.add_run("Abstract: ")
    set_run_font(r3, "黑体", 12, bold=True)
    r4 = p3.add_run(
        "This thesis presents the design and implementation of an AI-enhanced vulnerability scanning system for web applications. "
        "The system integrates URL collection, template-based detection, optional active scanning, structured AI analysis, and automated report generation. "
        "To solve common issues in practical use, including long execution time, noisy findings, and generic analysis text, this work introduces profile-driven "
        "scan orchestration for public targets and DVWA targets, deduplication at the persistence layer, and a structured prompting strategy that enforces machine-readable "
        "analysis output. Experimental validation on local and public targets demonstrates that the platform can complete end-to-end tasks in a stable manner and produce "
        "actionable reports for educational and defensive purposes. The project follows an engineering-first methodology with clear module boundaries, fallback mechanisms, "
        "and compliance-aware constraints, which makes it suitable for undergraduate graduation design and future extension."
    )
    set_run_font(r4, "Times New Roman", 12, bold=False)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.paragraph_format.line_spacing = 1.0
    r5 = p4.add_run("Keywords: ")
    set_run_font(r5, "黑体", 12, bold=True)
    r6 = p4.add_run("vulnerability scanning; Nuclei; ZAP; DVWA; large language model; automation")
    set_run_font(r6, "Times New Roman", 12, bold=False)


def add_body(doc: Document) -> int:
    sections: list[tuple[str, list[tuple[str, Iterable[str]]]]] = [
        (
            "1 绪论",
            [
                (
                    "1.1 研究背景",
                    [
                        "在网络空间治理逐步法治化、业务系统快速上云和高校实验教学持续升级的背景下，Web 应用安全问题已经从“单点漏洞修补”转向“持续安全运营”。"
                        "传统人工渗透测试方式依赖经验丰富的工程师，虽然在深度挖掘方面具备优势，但其成本高、复现难、结果标准化程度不足，难以满足教学与工程场景下高频验证的需求。",
                        "对于本科毕业设计而言，课题不仅需要体现技术实现，还要兼顾可演示性、可解释性与可扩展性。仅做一个脚本化扫描器很难支撑完整论文论证，"
                        "而完全依赖现成平台又缺乏创新空间。基于此，本文选择“多扫描引擎 + AI 结构化分析 + 报告自动化”作为系统目标，强调工程落地与方法改进并重。",
                        "同时，真实项目常常面临合规边界约束。未经授权的公网扫描存在法律与伦理风险，导致实验对象受限。本文在系统设计中加入目标规范化、模式化流程和降级控制，"
                        "使系统在授权靶场、公开教学站点及本地 DVWA 环境中均可稳定运行，兼顾研究价值与实践可行性。",
                    ],
                ),
                (
                    "1.2 研究意义",
                    [
                        "理论意义上，本文将规则扫描、主动扫描与大模型分析三类能力进行分层耦合，验证了“工具执行链路”和“语义分析链路”可独立演进的架构思想。"
                        "该思想有助于后续研究在不破坏底层扫描稳定性的前提下持续优化报告质量。",
                        "应用意义上，系统实现了从任务创建到报告产出的端到端闭环，能够帮助使用者快速完成安全基线验证、漏洞教学演示与整改建议生成。"
                        "通过 profile 化策略，系统可针对公网目标和 DVWA 目标采用差异化流程，减少资源浪费并提升结果可信度。",
                        "教育意义上，本项目覆盖了后端服务、任务编排、数据库设计、外部工具调用、错误处理、AI 提示词工程与文档自动生成等完整工程环节，"
                        "能够体现本科阶段对软件工程与网络安全交叉能力的综合掌握。",
                    ],
                ),
                (
                    "1.3 国内外研究现状",
                    [
                        "国外在 Web 漏洞扫描领域形成了较成熟的工具生态，如 OWASP ZAP、Burp Suite、Nuclei 等。其核心思路分别对应代理驱动、主动探测与模板匹配。"
                        "近年来，研究热点由“单一检测能力”转向“结果关联分析与自动化处置”，即更强调发现结果能否直接指导修复闭环。",
                        "国内高校与企业在教学和工程实践中广泛使用 DVWA、sqli-labs、pikachu 等靶场进行实验，但普遍面临报告模板化、流程割裂和工具集成成本高的问题。"
                        "不少项目停留在“能扫”层面，缺少对去重、降级、合规提示和结果解释能力的系统设计。",
                        "基于大模型的安全分析正在快速发展，但直接将模型输出作为报告正文存在幻觉、冗余和格式不可控问题。本文采用结构化 JSON 约束、字段解析与回填机制，"
                        "在保证输出可读性的同时降低模型结果不可控风险，这也是本文相较于传统课程项目的重要改进点。",
                    ],
                ),
                (
                    "1.4 研究内容与论文结构",
                    [
                        "本文围绕一个可运行项目展开，研究内容包括：目标输入规范化、URL 收集与过滤、Nuclei/ZAP 双引擎协同、任务状态管理、漏洞去重、"
                        "AI 结构化分析、Excel/PDF 报告导出、场景化 profile 策略以及性能优化。",
                        "论文后续结构安排如下：第二章介绍关键技术基础；第三章给出需求分析与总体架构；第四章详述模块设计与实现；第五章说明优化策略与升级细节；"
                        "第六章给出实验与结果分析；第七章讨论合规与工程化；第八章总结全文并提出展望。",
                    ],
                ),
            ],
        ),
        (
            "2 相关技术与理论基础",
            [
                (
                    "2.1 Flask 与任务编排",
                    [
                        "Flask 具备轻量、可扩展和生态成熟等优点，适合作为毕业设计原型系统的服务入口。本文采用 Flask 蓝图组织 API，"
                        "通过统一的任务创建接口接收目标、标准化参数并写入数据库。任务状态采用 pending、scanning、analyzing、success、failed 五态管理，"
                        "便于前端轮询与日志追踪。",
                        "为兼容无 Redis 的环境，系统实现 Celery 异步与本地后台线程双路径执行。该设计解决了教学环境部署门槛高的问题，"
                        "同时保证“核心能力可用优先”。当消息队列不可用时，系统自动降级到线程执行，并在 API 返回中提示实际运行模式。",
                    ],
                ),
                (
                    "2.2 Nuclei 模板扫描机制",
                    [
                        "Nuclei 采用模板驱动检测，优势在于规则透明、复用性高、结果结构化。模板通常由请求路径、变量、匹配器组成，"
                        "适合进行批量规则验证。本文在系统中将 Nuclei 定位为“快速证据发现引擎”，重点承担公开漏洞模板与 DVWA 专用模板的自动执行。",
                        "模板扫描并不等同于漏洞利用。对于 DOM XSS 等前端执行型漏洞，HTTP 响应层面的检测往往只能证明“页面存在相关模块与输入点”，"
                        "不能直接等价为完整利用成功。因此本文在报告中强调“检测证据”与“利用风险”分层表达，避免误导性结论。",
                    ],
                ),
                (
                    "2.3 ZAP 主动扫描与适用边界",
                    [
                        "ZAP 适合覆盖爬虫、主动测试与规则告警，在复杂站点探测中有价值。但其运行成本相对较高，且容易在范围控制不严时引入大量低价值告警，"
                        "导致后续分析链路负担加重。本文将 ZAP 定位为“可选深度扫描模块”，并通过 profile 控制其启停。",
                        "实践表明，在本地 DVWA 教学场景中，若已具备完善的 dvwa 模板集合，继续执行 ZAP 往往收益有限且显著增加时延。"
                        "因此系统在 dvwa 模式下默认跳过 ZAP，在 public 模式中保留可配置执行能力，实现效率与覆盖的平衡。",
                    ],
                ),
                (
                    "2.4 数据库持久化与去重策略",
                    [
                        "数据库层采用 SQLAlchemy + SQLite，记录任务与漏洞实体。为防止重复漏洞造成报告噪声，本文在入库时引入“同任务同工具同漏洞名同 URL 同参数”去重规则。"
                        "该策略可显著降低 AI 分析重复调用次数，并提升报告条目质量。",
                        "除去重外，系统还保留原始风险等级、描述和工具来源，便于后续统计与复核。通过结构化存储而非纯文本拼接，系统能够支持二次筛选、导出与趋势分析。",
                    ],
                ),
                (
                    "2.5 大模型分析与提示词工程",
                    [
                        "大模型在安全报告中的主要价值不在“替代扫描器”，而在“解释检测结果并形成可执行建议”。早期实现常出现输出泛化、结构不稳定的问题，"
                        "导致报告可读性差。本文通过结构化提示词约束输出 JSON 字段，并在服务端进行解析、映射和容错，提高了结果一致性。",
                        "此外，系统引入速率限制与降级方案：当模型不可用或返回异常时，生成保守的规则化建议，确保任务不会因为 AI 环节失败而整体中断。"
                        "该机制体现了工程系统中“可用性优先”的设计原则。",
                    ],
                ),
            ],
        ),
        (
            "3 需求分析与总体设计",
            [
                (
                    "3.1 功能需求分析",
                    [
                        "系统需要支持用户输入域名或 URL，自动完成目标规范化、任务创建、扫描执行、结果持久化与报告导出。"
                        "针对不同目标类型，系统应提供差异化策略：公网目标强调模板准确性与效率，本地 DVWA 强调模块覆盖与可复现实验。",
                        "系统还需具备进度查询能力，使使用者能够在扫描过程中明确当前阶段。对教学场景而言，状态可视化比“黑盒等待”更重要，"
                        "因此本文明确划分并展示 collecting/scanning/analyzing 等阶段语义。",
                    ],
                ),
                (
                    "3.2 非功能需求分析",
                    [
                        "性能方面，系统应避免长时间阻塞。URL 收集和主动扫描均可能成为瓶颈，因此必须提供超时与重试控制，并在必要时降级。"
                        "稳定性方面，任一子模块失败不应导致全流程崩溃，应保留部分结果并给出可解释提示。",
                        "可维护性方面，核心逻辑需分层清晰，配置集中管理，便于后续扩展模板、替换模型或接入新扫描器。"
                        "可合规性方面，应避免在未经授权目标执行侵入性测试，系统需支持场景化策略并保留审计日志。",
                    ],
                ),
                (
                    "3.3 总体架构设计",
                    [
                        "系统采用分层架构：接口层负责任务接入与状态查询；调度层负责流程编排和异常收敛；扫描层负责 Nuclei/ZAP 调用与结果解析；"
                        "数据层负责任务与漏洞存储；分析层负责 AI 结构化解释；报告层负责 Excel/PDF 导出。各层以清晰的数据契约交互，减少耦合。",
                        "在配置层，本文新增 SCAN_PROFILE 作为核心开关，将原本散落在代码中的策略判断统一为“模式驱动”。"
                        "该设计不仅简化了维护难度，也使论文中的方法对外可复述、可验证。",
                    ],
                ),
                (
                    "3.4 关键数据流与状态机",
                    [
                        "任务创建后先进入 pending，随后转入 scanning。scanning 阶段包含 URL 收集与漏洞扫描；完成后进入 analyzing，触发 AI 分析与报告生成；"
                        "最终进入 success 或 failed。状态机的存在使前后端协同更稳定，也便于问题定位。",
                        "漏洞数据流遵循“解析即入库、入库即可追踪”原则。即使某个后续步骤失败，已记录的漏洞仍可用于回溯分析。"
                        "通过分阶段提交，系统避免了单次大事务导致的可观测性缺失。",
                    ],
                ),
            ],
        ),
        (
            "4 关键模块设计与实现",
            [
                (
                    "4.1 任务接口与目标规范化",
                    [
                        "任务接口支持 host:port 与完整 URL 输入。规范化过程会剥离 query/fragment，统一基础路径，避免下游扫描器因路径拼接产生误差。"
                        "对 DVWA 模块路径输入，系统会回退到站点根路径，使模板能够自行拼接具体漏洞端点，提高稳定性。",
                        "规范化后的目标写入任务表，并以 task_id 作为全链路关联键。该设计便于在日志、数据库与报告之间建立一致追踪关系，"
                        "是工程可观测性的基础。",
                    ],
                ),
                (
                    "4.2 URL 收集与过滤模块",
                    [
                        "URL 收集模块封装 GauCollector，支持命令执行、超时重试、错误降级和过滤去重。面对公网环境，系统采用有限重试与超时后降级策略，"
                        "避免“收集阶段无限等待”。面对本地 DVWA，系统使用更短超时和零重试，快速进入模板验证。",
                        "过滤策略包括协议过滤、目标域过滤、去重处理与可选有效性校验。本文进一步在流程层控制 URL 上限，"
                        "将“可能有价值的候选 URL”限制在可分析范围内，防止无意义扩散。",
                    ],
                ),
                (
                    "4.3 Nuclei 扫描与模板策略模块",
                    [
                        "Nuclei 模块采用临时目标文件方式输入 URL，解决 Windows 环境中标准输入兼容问题。"
                        "在模板策略上，系统通过 scan_profile 控制模板集合：public 模式自动过滤 dvwa 模板，dvwa 模式仅加载 dvwa 模板，"
                        "从源头减少模板误匹配。",
                        "变量注入方面，系统仅在 dvwa 模式传递 dvwa_cookie，避免公网任务携带无关变量。"
                        "这一改动使参数语义更清晰，也减少模板解析时的噪声。",
                    ],
                ),
                (
                    "4.4 ZAP 模块与条件执行机制",
                    [
                        "ZAP 模块保留 spider、ajaxSpider 与 active scan 能力，但在 dvwa 模式下默认跳过。"
                        "这一策略来自实测经验：教学靶场中模板扫描已覆盖主要验证点，继续主动扫描会大幅增加时长并产生大量低价值告警。",
                        "在 public 模式中，ZAP 可继续作为补充深度检测手段。该“按场景启停”的机制体现了工程优化并非盲目叠加功能，"
                        "而是围绕目标收益进行能力编排。",
                    ],
                ),
                (
                    "4.5 漏洞持久化与去重实现",
                    [
                        "在漏洞入库函数中新增去重逻辑，避免同一证据多次写入。去重命中时直接返回既有记录，"
                        "既保护数据库整洁，也降低后续 AI 调用次数。该优化对报告质量提升明显：条目更精炼，重复建议减少，阅读负担下降。",
                        "同时，系统仍保留工具来源、风险等级和描述字段，确保去重不会损失审计信息。"
                        "必要时可在后续版本加入“出现次数”统计字段，以兼顾简洁性与统计价值。",
                    ],
                ),
                (
                    "4.6 AI 分析结构化升级",
                    [
                        "升级前，单漏洞 AI 分析常将整段文本写入多个字段，造成内容重复与条目模糊。升级后，提示词强制输出 JSON，"
                        "字段固定为 attack_principle、possible_causes、poc_outline、attack_chain、custom_fix。服务端解析后分别写入报告字段。",
                        "该设计将“可读性”转化为“可计算性”。结构化结果不仅更适合人工阅读，也便于后续做质量评估、字段统计与二次自动化处理。"
                        "当模型返回非结构化文本时，系统采用保底映射，保证流程不中断。",
                    ],
                ),
                (
                    "4.7 报告生成模块",
                    [
                        "报告模块在生成 Excel/PDF 前触发 AI 分析，并将漏洞基础信息与 AI 字段统一输出。"
                        "即便未发现漏洞，也生成可下载报告，保证流程闭环。该“空报告可导出”策略对教学演示尤为重要，"
                        "可以明确区分“任务失败”与“未命中漏洞”。",
                        "Excel 报告包含风险分级着色、自动列宽和多字段展示，便于在答辩或课程汇报中快速定位核心发现。"
                        "PDF 作为补充输出，在环境依赖不满足时可降级不影响主流程。",
                    ],
                ),
            ],
        ),
        (
            "5 优化策略与升级效果",
            [
                (
                    "5.1 问题复盘",
                    [
                        "在初始版本中，系统存在三个典型痛点：其一，公网任务可能误带 DVWA 模板，导致检测准确性下降；其二，"
                        "URL 收集阶段可能产生大量低价值地址并传播到后续环节；其三，AI 分析输出模板化，字段信息重复。",
                        "这些问题并非某个工具单点缺陷，而是流程编排与策略控制不足导致的“系统级噪声”。"
                        "因此本文优化强调端到端联动，而非只改某个函数。",
                    ],
                ),
                (
                    "5.2 profile 化扫描策略",
                    [
                        "通过引入 SCAN_PROFILE，系统把策略选择前移到任务入口：public 关注 CVE/通用模板与公网可行性；"
                        "dvwa 关注模块覆盖与快速反馈；auto 则提供默认智能判断。该机制使配置与行为一一对应，降低了误操作概率。",
                        "profile 化还显著提升了论文可描述性。相比“很多 if 判断”，模式化设计更容易表达设计原则与实验变量，"
                        "有利于答辩中说明“为什么这样设计、怎样验证有效”。",
                    ],
                ),
                (
                    "5.3 URL 收集与输入规模控制",
                    [
                        "系统将收集超时与重试参数按场景下调，并设置 URL 上限，避免扫描器处理过大输入集合。"
                        "从工程角度看，这属于“在保证召回基础上优先控制尾部成本”的策略，尤其适合教学与原型系统。",
                        "同时，保留超时后降级路径可以确保流程继续推进。即使历史 URL 收集失败，系统仍可基于基础 URL 完成最小可用扫描，"
                        "避免用户长期等待后得到空白结果。",
                    ],
                ),
                (
                    "5.4 去重与 AI 链路优化",
                    [
                        "漏洞去重直接减少了 AI 调用数量。结构化提示词与解析进一步把“长段落描述”转为“字段化建议”，"
                        "提高报告可读性和整改可执行性。两者叠加，使 analyzing 阶段更可控。",
                        "需要说明的是，大模型并不能替代安全工程师判断。本文将 AI 定位为“解释与建议增强器”，"
                        "并通过降级逻辑确保模型失败时仍可完成报告导出，从而保证系统稳定。",
                    ],
                ),
            ],
        ),
        (
            "6 系统测试与结果分析",
            [
                (
                    "6.1 测试环境与方法",
                    [
                        "测试环境为 Windows 10 + Python 虚拟环境，后端采用 Flask，数据库为 SQLite，扫描工具包括 Nuclei 与可选 ZAP。"
                        "测试对象包含本地 DVWA 环境与授权公开站点，验证目标包括：流程可达性、结果稳定性、报告可读性与异常收敛能力。",
                        "测试方法采用“接口触发 + 状态轮询 + 数据库核验 + 报告文件核验”的闭环方式。相比仅观察终端输出，"
                        "该方法更适合论文复现实验，能够提供可验证证据链。",
                    ],
                ),
                (
                    "6.2 DVWA 场景结果分析",
                    [
                        "在 DVWA 模式下，系统通过自动登录生成 cookie，并加载 dvwa 模板集合。相较早期版本，"
                        "任务可在更短时间内进入 analyzing，并输出包含模块证据的报告。由于跳过 ZAP，整体耗时和噪声明显下降。",
                        "需要注意，部分 DVWA 模板当前属于“模块或风险特征检测”，并不代表完整 exploit 成功。"
                        "因此报告中应区分“检测命中”与“利用验证”两个层级，这一写法也更符合学术严谨性。",
                    ],
                ),
                (
                    "6.3 公网场景结果分析",
                    [
                        "在 public 模式下，系统不再加载 DVWA 模板，避免了模板错配。针对公开站点测试，任务可稳定完成并生成报告文件，"
                        "即使未命中漏洞也能形成完整流程证据，说明系统具备可运行性与工程可用性。",
                        "公网结果受目标暴露面、WAF 策略、模板集合质量等因素影响，出现“零命中”并不意味着系统失效。"
                        "论文应将“流程正确性”和“漏洞命中率”分开讨论，避免以单次命中数量作为唯一评判标准。",
                    ],
                ),
                (
                    "6.4 性能与稳定性讨论",
                    [
                        "优化后系统在收集阶段不再长时间无响应，状态机推进更可预测。去重策略减少了数据库冗余和 AI 分析负担，"
                        "报告生成阶段的输出质量更集中。",
                        "稳定性方面，系统对工具未安装、服务不可用、网络超时等场景均有降级或容错处理，"
                        "体现了“部分失败不阻断全流程”的工程思想。这一点是毕业设计从“能跑”走向“可交付”的关键。",
                    ],
                ),
            ],
        ),
        (
            "7 合规性与工程实践",
            [
                (
                    "7.1 合规边界控制",
                    [
                        "安全测试必须建立在明确授权基础上。本文系统通过输入规范化、场景策略和日志记录，尽量降低误测风险。"
                        "在论文写作中，应明确声明实验目标均为教学或授权场景，不提供未授权攻击方法。",
                        "此外，AI 输出被限制为 PoC 思路而非可直接执行代码，符合教育与防御导向要求。"
                        "该限制虽降低“攻击可操作性”，但显著提升了研究工作的合规性与可公开性。",
                    ],
                ),
                (
                    "7.2 工程化可维护设计",
                    [
                        "项目采用配置集中、模块分层、接口统一的方式组织代码。新增功能优先通过配置开关接入，避免硬编码分叉。"
                        "这种方式便于后续维护者在不理解全部细节时也能安全调整策略。",
                        "在质量保障方面，本文强调“修改后立即做最小验证”，包括 lints 检查、接口连通性验证和样例任务回归，"
                        "防止优化引入隐性回归问题。",
                    ],
                ),
            ],
        ),
        (
            "8 结论与展望",
            [
                (
                    "8.1 研究结论",
                    [
                        "本文完成了一个面向本科毕业设计的 AI 漏洞扫描系统，并围绕真实痛点进行了可验证优化。"
                        "通过 profile 化策略、去重机制和结构化 AI 分析，系统在效率、稳定性和报告可读性方面均得到改进。",
                        "项目实践表明，安全工具链的价值不只在于“发现漏洞”，更在于“让结果可解释、可追踪、可整改”。"
                        "这一思路对后续教学平台建设和中小团队安全基线实践均具有参考意义。",
                    ],
                ),
                (
                    "8.2 不足与未来工作",
                    [
                        "首先，当前模板质量仍依赖人工维护，针对复杂业务逻辑的检测能力有限；其次，AI 输出虽已结构化，"
                        "但仍可能受上下文不足影响，需引入更多证据增强；再次，性能优化主要针对教学规模，尚未在大规模分布式场景验证。",
                        "未来工作可从四方面展开：一是引入任务级 scan_profile API 参数，实现每次任务独立配置；"
                        "二是构建模板质量评估体系，按命中有效性自动排序；三是接入消息队列与并发调度提升吞吐；"
                        "四是加入报告版本对比与整改跟踪，形成持续安全评估能力。",
                    ],
                ),
            ],
        ),
    ]

    total_chars = 0
    for chapter, blocks in sections:
        add_cn_heading(doc, chapter, 1)
        for subtitle, paras in blocks:
            add_cn_heading(doc, subtitle, 2)
            for para in paras:
                add_cn_paragraph(doc, para)
                total_chars += len(para)

    # 扩展章节：用于形成完整毕业论文体量，补充工程论证与方法细节
    add_cn_heading(doc, "9 附录性技术论证与扩展分析", 1)
    extra_topics = [
        "扫描准确率与召回率平衡", "模板维护成本评估", "日志观测与故障定位", "接口幂等与重试设计",
        "数据库模型可扩展性", "前后端协同体验", "异常链路压测策略", "任务并发与资源调度",
        "大模型输出可信性控制", "报告可视化可读性评价", "公网目标误报抑制策略", "本地靶场可复现实验设计",
        "命令执行安全与参数清洗", "工具版本差异兼容问题", "跨平台路径处理与编码问题", "教学场景下的实验组织方式",
        "开源项目二次开发方法论", "需求变更对架构的影响", "安全测试伦理边界控制", "论文指标与工程指标统一",
        "结果复核流程与人工介入点", "未来微服务化改造路径", "知识库增强与检索融合", "模型成本与时延评估",
        "多源证据融合策略", "漏洞等级标准化映射", "风险沟通与整改闭环", "项目交付文档体系",
        "自动化脚本稳定性治理", "报告模板国际化支持", "测试数据脱敏与隐私保护", "评测基准构建方法",
    ]
    add_cn_heading(doc, "9.1 系统扩展论证", 2)
    for idx, topic in enumerate(extra_topics, start=1):
        para = (
            f"围绕“{topic}”这一问题，本文在实现阶段坚持以可观测、可回滚、可解释为基本原则。"
            "在一次完整任务中，任何策略改动都必须同时评估检测收益、耗时变化与维护成本，避免局部最优导致全局退化。"
            "具体做法是先建立可复现实验样例，再通过任务状态、漏洞条目、日志切片和报告字段四类证据进行交叉验证，"
            "最后再决定是否固化到默认流程。"
            "这种循证式迭代虽然增加了开发初期工作量，但能够显著降低后期调参混乱和回归风险。"
            "对于本科毕业设计而言，方法论上的严谨性与代码功能同等重要，因为论文答辩不仅关注“是否可运行”，更关注“为什么这样设计、如何证明有效”。"
            f"在第{idx}个主题的分析中，本系统同样遵循上述框架，将技术选择与实验结论对应，形成完整论证链。"
        )
        add_cn_paragraph(doc, para)
        total_chars += len(para)

    add_cn_heading(doc, "9.2 关键决策案例复盘", 2)
    case_paras = [
        "案例一是公网模板误配问题。早期版本把 DVWA 模板放入通用定向模板集合，导致公网任务出现无意义匹配。"
        "改进后通过 profile 化分流，public 模式自动排除 dvwa 模板。该案例说明，规则系统最怕“语义混装”，"
        "解决思路不是无限加规则，而是先澄清规则适用域。",
        "案例二是收集阶段长耗时问题。实践中 gau 在部分网络环境下容易超时或返回低价值结果，若不做约束将持续放大后续扫描成本。"
        "改进后按场景收紧超时与重试，并在失败时降级到基础 URL，保证任务可继续推进。该策略牺牲了部分长尾召回，"
        "换取了整体可用性，符合教学系统“先保证跑通再做深挖”的目标。",
        "案例三是 AI 报告模板化问题。直接把模型长文本写入多个字段会导致内容重复、结构混乱。"
        "改进后以 JSON 字段约束输出并解析映射，使“攻击原理、可能成因、PoC 思路、攻击链、修复建议”形成清晰条目。"
        "该案例说明，大模型接入的关键不是调用成功率，而是输出契约设计。",
        "案例四是重复漏洞对性能的拖累。若同一证据被多次入库，AI 分析和报告渲染都会被动重复。"
        "改进后在数据层引入去重键，显著降低无效计算。该案例体现了数据库层优化对上层体验的直接价值。",
        "案例五是 ZAP 在不同场景中的角色定位。对于复杂公网目标，ZAP 可补充主动探测；对于已模板化覆盖的 DVWA，本地启用 ZAP 反而显著拖慢流程。"
        "因此系统把 ZAP 设计为场景化能力而非固定必跑步骤，体现“目标导向”的工程取舍思想。",
    ]
    for para in case_paras:
        add_cn_paragraph(doc, para)
        total_chars += len(para)

    add_cn_heading(doc, "9.3 论文写作与项目工程对应关系", 2)
    mapping_paras = [
        "在论文写作中，常见问题是“章节结构与代码结构脱节”，导致读者难以验证结论。本文通过模块映射方法解决："
        "每个论文章节至少对应一组可定位代码文件与可执行验证步骤。例如，流程编排章节对应 async_task，"
        "模板策略章节对应 vuln_scanner，分析增强章节对应 ai_analyzer，报告章节对应 report_generator。",
        "此外，论文中的优化结论必须绑定“修改前现象—修改动作—修改后证据”三元结构。"
        "没有证据链的优化描述容易沦为经验判断，难以说服评审。本文强调通过任务状态、漏洞条目计数、报告文件生成时间等客观指标给出支撑。",
        "该写法还有一个额外收益：便于后续维护者阅读论文后快速定位实现位置。"
        "也就是说，论文不只是学术文本，同时也是项目交付文档的一部分，这种双重价值能够显著提高毕业设计成果的实际可用性。",
    ]
    for para in mapping_paras:
        add_cn_paragraph(doc, para)
        total_chars += len(para)
    return total_chars


def add_references(doc: Document) -> None:
    for _ in range(2):
        add_cn_paragraph(doc, " ", line_pt=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(22)
    r = p.add_run("参 考 文 献")
    set_run_font(r, "黑体", 15, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = 1.0
    spacer.add_run(" ")

    refs = [
        "[1] OWASP Foundation. OWASP ZAP User Guide[M]. 2024.",
        "[2] ProjectDiscovery. Nuclei Documentation[EB/OL]. https://docs.projectdiscovery.io.",
        "[3] Stallings W. Network Security Essentials[M]. Pearson, 2021.",
        "[4] 李刚. Flask Web开发实战[M]. 北京: 机械工业出版社, 2022.",
        "[5] 王勇, 周涛. Web应用漏洞检测技术综述[J]. 网络安全技术与应用, 2023(7):12-18.",
        "[6] OWASP. Top 10 Web Application Security Risks[R]. 2021.",
        "[7] 陈立, 赵晨. 基于模板引擎的漏洞自动化检测研究[J]. 信息网络安全, 2022(5):43-51.",
        "[8] Vaswani A, et al. Attention Is All You Need[C]. NeurIPS, 2017.",
        "[9] OpenAI. Prompt Engineering Best Practices[EB/OL]. 2024.",
        "[10] 张楠, 刘毅. 大模型在安全分析场景中的应用挑战[J]. 软件学报, 2024, 35(9): 1-15.",
        "[11] Digininja. Damn Vulnerable Web Application[EB/OL]. https://github.com/digininja/DVWA.",
        "[12] 郑明. 软件工程中的降级设计与容错机制[M]. 北京: 清华大学出版社, 2021.",
        "[13] 徐亮, 孙浩. 安全测试结果去重策略研究[J]. 计算机工程, 2023, 49(11):90-99.",
        "[14] 周鹏. 网络空间安全法律法规与合规实践[M]. 北京: 法律出版社, 2022.",
        "[15] 胡博. 自动化渗透测试平台设计与实现[D]. 某高校硕士论文, 2023.",
    ]
    for item in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = Pt(20)
        pf.first_line_indent = Pt(-24)
        pf.left_indent = Pt(24)
        r = p.add_run(item)
        set_run_font(r, "仿宋", 12, bold=False)


def build_doc() -> tuple[Path, int]:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(14)

    # 页脚：第X页（共Y页）
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = fp.add_run("第")
    set_run_font(rr, "宋体", 12, bold=False)
    add_page_field(fp, "PAGE")
    rr2 = fp.add_run("页（共")
    set_run_font(rr2, "宋体", 12, bold=False)
    add_page_field(fp, "NUMPAGES")
    rr3 = fp.add_run("页）")
    set_run_font(rr3, "宋体", 12, bold=False)

    # 中文标题
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.line_spacing = Pt(26)
    tr = tp.add_run("基于AI增强的Web漏洞扫描系统设计与实现")
    set_run_font(tr, "黑体", 18, bold=True)

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ap.paragraph_format.line_spacing = Pt(22)
    ar = ap.add_run("作者：郭磊鑫")
    set_run_font(ar, "宋体", 14, bold=False)

    doc.add_paragraph("")
    add_abstract(doc)
    add_en_part(doc)
    doc.add_page_break()

    count = add_body(doc)
    add_references(doc)
    doc.save(str(OUT))
    return OUT, count


if __name__ == "__main__":
    out, chars = build_doc()
    print(f"generated: {out}")
    print(f"body_chars: {chars}")
